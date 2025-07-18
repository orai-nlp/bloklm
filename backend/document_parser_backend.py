import os
import io
# from werkzeug.datastructures import FileStorage
import PyPDF2
import docx
import chardet
import textract


def extract_text_from_document(file) -> dict:
    """
    Extract text from various document types (PDF, DOC/DOCX, TXT, SRT)
    
    Args:
        file: FileStorage object from Flask request.files
        
    Returns:
        dict: {
            'success': bool,
            'text': str,
            'filename': str,
            'file_type': str,
            'error': str (if success=False)
        }
    """
    
    # ERROR
    if not file or not file.filename:
        return {
            'success': False,
            'text': '',
            'filename': '',
            'file_type': '',
            'error': 'No file provided'
        }
    
    # INIT
    filename = file.filename
    file_extension = os.path.splitext(filename)[1].lower()
    
    try:
        # Read file content into memory
        file_content = file.read()
        file.seek(0)  # Reset file pointer for potential reuse
        
        # Determine file type and extract text
        if file_extension == '.pdf':
            text = _extract_pdf_text(file_content)
            file_type = 'PDF'
            
        elif file_extension == '.docx':
            text = _extract_docx_text(file_content)
            file_type = 'DOCX'

        elif file_extension == '.doc':
            text = _extract_doc_text(file_content)
            file_type = 'DOC'
            
        elif file_extension == '.txt':
            text = _extract_txt_text(file_content)
            file_type = 'TXT'
            
        elif file_extension == '.srt':
            text = _extract_srt_text(file_content)
            file_type = 'SRT'
            
        else:
            return {
                'success': False,
                'text': '',
                'filename': filename,
                'file_type': file_extension,
                'error': f'Unsupported file type: {file_extension}'
            }
        
        return {
            'success': True,
            'text': text,
            'filename': filename,
            'file_type': file_type,
            'error': ''
        }
        
    except Exception as e:
        return {
            'success': False,
            'text': '',
            'filename': filename,
            'file_type': file_extension,
            'error': f'Error processing file: {str(e)}'
        }

# PDF formatutik testua lortzeko
def _extract_pdf_text(file_content: bytes) -> str:
    """Extract text from PDF file"""
    text = ""
    try:
        pdf_file = io.BytesIO(file_content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            text += page.extract_text() + "\n"
            
    except Exception as e:
        raise Exception(f"Error extracting PDF text: {str(e)}")
    
    return text.strip()

# DOCX formatutik testua lortzeko
def _extract_docx_text(file_content: bytes) -> str:
    """Extract text from DOCX file"""
    try:
        doc_file = io.BytesIO(file_content)
        doc = docx.Document(doc_file)
        
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
            
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\t"
                text += "\n"
                
    except Exception as e:
        raise Exception(f"Error extracting DOCX text: {str(e)}")
    
    return text.strip()

# DOC formatutik testua lortzeko
def _extract_doc_text(file_content: bytes) -> str:
    """Extract text from DOC file using textract from in-memory bytes"""
    try:
        with tempfile.NamedTemporaryFile(delete=True, suffix=".doc") as tmp:
            tmp.write(file_content)
            tmp.flush()  # Ensure all data is written
            text = textract.process(tmp.name).decode('utf-8')
            return text.strip()
    except Exception as e:
        raise Exception(f"Error extracting DOC text from bytes: {str(e)}")

# TXT formatutik testua lortzeko, encoding kasuak kasu
def _extract_txt_text(file_content: bytes) -> str:
    """Extract text from TXT file with encoding detection"""
    try:
        # Detect encoding
        detected = chardet.detect(file_content)
        encoding = detected['encoding'] or 'utf-8'
        
        # Decode with detected encoding
        text = file_content.decode(encoding)
        
    except Exception as e:
        # Fallback to common encodings
        for encoding in ['utf-8', 'latin-1', 'cp1252']:
            try:
                text = file_content.decode(encoding)
                break
            except:
                continue
        else:
            raise Exception(f"Could not decode text file: {str(e)}")
    
    return text.strip()

# SRT formatutik testua lortzeko
def _extract_srt_text(file_content: bytes) -> str:
    """Extract text from SRT subtitle file"""
    try:
        # Detect encoding
        detected = chardet.detect(file_content)
        encoding = detected['encoding'] or 'utf-8'
        
        # Decode content
        content = file_content.decode(encoding)
        
        # Parse SRT format
        text = ""
        lines = content.strip().split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i].strip()
            
            # Skip empty lines
            if not line:
                i += 1
                continue
            
            # Skip subtitle numbers (digits only)
            if line.isdigit():
                i += 1
                continue
            
            # Skip timestamp lines (contain -->)
            if '-->' in line:
                i += 1
                continue
            
            # This should be subtitle text
            text += line + " "
            i += 1
            
    except Exception as e:
        raise Exception(f"Error extracting SRT text: {str(e)}")
    
    return text.strip()


# Requirements to install:
# pip install PyPDF2 python-docx chardet flask

# Flask usage
"""

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    result = extract_text_from_document(file)
    
        
    if result['success']:
        return jsonify({
            'success': True,
            'text': result['text'],
            'filename': result['filename'],
            'file_type': result['file_type']
        })
    else:
        return jsonify({
            'success': False,
            'error': result['error']
        }), 400

"""